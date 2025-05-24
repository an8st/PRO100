import io
import json
import uuid

import ffmpeg
from docx import Document
from flask import Blueprint, render_template, request, jsonify, send_file, request
from vosk import Model, KaldiRecognizer
import wave
import os
import re

from app import socketio
from app.model.Table import Table

main_bp = Blueprint('main', __name__)
tables = {}
FFMPEG_BIN = r'C:\ProgramData\chocolatey\bin\ffmpeg.exe'

def new_table_no_params(name: str):
    new_table(name, 10, 10)
    return open_table(name)
def new_table(name: str, rows: int, columns: int):
    if name in tables:
        raise ValueError(f"Table '{name}' already exists")
    col_names = [f"{i+1}" for i in range(columns)]
    table = Table(name, col_names)
    for _ in range(rows):
        table.add_row(["" for _ in range(columns)])
    tables[name] = table
    return open_table(name)


def drop_table(name: str):
    if name not in tables:
        raise ValueError(f"Table '{name}' does not exist")
    del tables[name]


def update_table(name: str, row: int, col: int, data: str):
    if name not in tables:
        raise ValueError(f"Table '{name}' does not exist")
    tables[name].update_cell(row, col, data)
    return open_table(name)


def open_table(name: str) -> str:
    print(name)
    print(tables)
    if name not in tables:
        raise ValueError(f"Table '{name}' does not exist")
    table = tables[name]
    return json.dumps(table.to_dict(), ensure_ascii=False, indent=2)


def recognize_speech_vosk(audio_file_path, model_path='path_to_your_vosk_model'):
    wf = wave.open(audio_file_path, "rb")
    
    model = Model(model_path)
    rec = KaldiRecognizer(model, wf.getframerate())
    
    full_text = ""
    while True:
        data = wf.readframes(4000)
        if len(data) == 0:
            break
        if rec.AcceptWaveform(data):
            result = json.loads(rec.Result())
            full_text += result.get("text", "") + " "
    
    final_result = json.loads(rec.FinalResult())
    full_text += final_result.get("text", "")
    print(full_text.strip())
    return full_text.strip()


# Обработчик POST /new_command
@main_bp.route('/new_command', methods=['POST'])
def new_command():
    if 'command' not in request.files:
        return jsonify({"error": "Audio not found"}), 400

    audio_file = request.files['command']

    temp_input = f"temp_{uuid.uuid4().hex}.webm"
    temp_output = f"temp_{uuid.uuid4().hex}.wav"

    audio_file.save(temp_input)

    try:
        ffmpeg.input(temp_input).output(temp_output).run(
            cmd=FFMPEG_BIN,
            quiet=True,
            overwrite_output=True
        )
        command_text = recognize_speech_vosk(temp_output, model_path='vosk-model-small-ru-0.22')
        print("Распознанная команда:", command_text)
    except Exception as e:
        return jsonify({"error": f"Recognition error: {e}"}), 500
    finally:
        if os.path.exists(temp_input): os.remove(temp_input)
        if os.path.exists(temp_output): os.remove(temp_output)

    
    def word_to_number(word):
        num_words = {
            'один': 1, 'одна': 1, 'одной': 1,
            'два': 2, 'две': 2, 'двумя': 2,
            'три': 3, 'тремя': 3,
            'четыре': 4, 'четырьмя': 4,
            'пять': 5, 'пятью': 5,
            'шесть': 6, 'шестью': 6,
            'семь': 7, 'семью': 7,
            'восемь': 8, 'восемью': 8,
            'девять': 9, 'девятью': 9,
            'десять': 10, 'десятью': 10
        }
        if word.isdigit():
            return int(word)
        return num_words.get(word.lower())

    patterns = [
        {
            'regex': r'добавить(?:\s+е|ть)?\s+таблицу\s+(\w+)',
            'func': new_table_no_params,
            'args': ['name']
        },
        {
            'regex': r'создать(?:\s+е|ть)?\s+таблицу\s+(\w+)',
            'func': new_table_no_params,
            'args': ['name']
        },
        {
            'regex': r'создать(?:\s+е|ть)?\s+таблицу\s+(\w+)\s+с\s+(\w+)\s+(?:строками|строка)\s*и\s*(\w+)\s+(?:колонками|колонка)',
            'func': new_table,
            'args': ['name', 'rows', 'columns']
        },
        {
            'regex': r'удалить(?:\s+е|ть)?\s+таблицу\s+(\w+)',
            'func': drop_table,
            'args': ['name']
        },
        {
            'regex': r'обновить(?:\s+е|ть)?\s+таблицу\s+(\w+)\s+на\s+строке\s+(\w+)\s+и\s+колонки\s+(\w+)\s+данными\s+(.+)',
            'func': update_table,
            'args': ['name', 'row', 'col', 'data']
        },
        {
            'regex': r'обновить(?:\s+е|ть)?\s+таблицу\s+(\w+)\s+на\s+строке\s+(\w+)\s+и\s+колонке\s+(\w+)\s+данными\s+(.+)',
            'func': update_table,
            'args': ['name', 'row', 'col', 'data']
        },
        {
            'regex': r'открыть(?:\s+е|ть)?\s+таблицу\s+(\w+)',
            'func': open_table,
            'args': ['name']
        }
    ]

    for pattern in patterns:
        match = re.match(pattern['regex'], command_text, re.IGNORECASE)
        if match:
            params = match.groups()
            kwargs = {}
            for i, arg_name in enumerate(pattern['args']):
                value = params[i]
                if arg_name in ['rows', 'columns', 'row', 'col']:
                    value = word_to_number(value)
                kwargs[arg_name] = value
            result = pattern['func'](**kwargs)
            return jsonify({"result": result})

    return jsonify({"error": "The command is not recognized or is not supported."}), 400


@main_bp.route('/')
def index():
    return render_template('index.html')

@main_bp.route('/tables', methods=['GET'])
def list_tables():
    all_tables = [table.to_dict() for table in tables.values()]
    return render_template('tables.html', tables=all_tables)

@main_bp.route('/tables/<name>/download', methods=['GET'])
def download_table(name):
    if name not in tables:
        return jsonify({"error": "Table not found"}), 404

    table = tables[name]
    doc = Document()

    # Заголовок
    doc.add_heading(table.name, level=1)

    # Подзаголовок
    if table.description:
        doc.add_paragraph(table.description)

    doc.add_paragraph("")  # пустая строка

    # Таблица
    if table.columns:
        num_cols = len(table.columns)
        num_rows = len(table.rows)

        doc_table = doc.add_table(rows=1 + num_rows, cols=num_cols)
        doc_table.style = 'Table Grid'

        # Заголовки
        hdr_cells = doc_table.rows[0].cells
        for idx, col_name in enumerate(table.columns):
            hdr_cells[idx].text = col_name

        # Данные
        for i, row in enumerate(table.rows):
            row_cells = doc_table.rows[i + 1].cells
            for j, cell in enumerate(row):
                row_cells[j].text = str(cell)

    # Отправка документа
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    return send_file(
        file_stream,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True,
        download_name=f"{table.name}.docx"
    )

@socketio.on('update_table')
def handle_update_table(data):
    print("ASDSDASDFASDDAASDASD")
    try:
        name = data['name']
        row = int(data['row'])
        col = int(data['col'])
        value = data['value']
        client_id = request.sid  # Получаем ID текущей сессии
        
        update_table(name, row, col, value)
        table_json = open_table(name)
        # Отправляем данные только конкретному клиенту
        socketio.emit('table_updated', table_json, room=client_id)
    except Exception as e:
        socketio.emit('table_updated', {"error": str(e)}, room=client_id)

new_table("один", 2 ,2)
